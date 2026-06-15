"""Convert Router - PDF/Image conversion endpoints"""
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse
import shutil
import uuid
import zipfile
import io
from pathlib import Path
from backend.services import ImageService, PDFService
import fitz
from docx import Document

router = APIRouter(prefix="/api/convert", tags=["convert"])

BASE = Path(__file__).parent.parent.parent
UPLOADS = BASE / "uploads"
OUTPUTS = BASE / "outputs"

def uid():
    return str(uuid.uuid4())

def cleanup(*paths):
    for p in paths:
        try:
            if p and Path(p).exists():
                if Path(p).is_dir():
                    shutil.rmtree(p)
                else:
                    Path(p).unlink()
        except:
            pass

@router.post("/pdf-to-image")
async def pdf_to_image(
    file: UploadFile = File(...),
    format: str = Form("png"),
    dpi: int = Form(150)
):
    """Convert PDF to images (PNG/JPG)"""
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files allowed")
    
    if format not in ["png", "jpg"]:
        raise HTTPException(status_code=400, detail="Format must be PNG or JPG")
    
    upload_id = uid()
    input_path = UPLOADS / f"{upload_id}.pdf"
    output_dir = OUTPUTS / upload_id
    
    try:
        # Save uploaded file
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)
        
        # Convert PDF to images
        images = ImageService.pdf_to_images(str(input_path), format, dpi)
        
        # Save images
        output_dir.mkdir(exist_ok=True)
        zip_path = OUTPUTS / f"{upload_id}_images.zip"
        
        with zipfile.ZipFile(zip_path, 'w') as zf:
            for idx, (filename, img_data) in enumerate(images):
                zf.writestr(f"{idx + 1:03d}_{filename}", img_data)
        
        return {
            "status": "success",
            "file_id": upload_id,
            "page_count": len(images),
            "format": format
        }
    except Exception as e:
        cleanup(input_path, output_dir)
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

@router.post("/image-to-pdf")
async def image_to_pdf(files: list[UploadFile] = File(...)):
    """Convert images to PDF"""
    if len(files) == 0:
        raise HTTPException(status_code=400, detail="At least one image required")
    
    upload_id = uid()
    image_paths = []
    output_path = OUTPUTS / f"{upload_id}_from_images.pdf"
    
    try:
        # Save all uploaded images
        for file in files:
            if not any(file.filename.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg']):
                raise HTTPException(status_code=400, detail="Only PNG and JPG images allowed")
            
            file_id = uid()
            img_path = UPLOADS / f"{file_id}.{file.filename.split('.')[-1]}"
            content = await file.read()
            
            with open(img_path, "wb") as f:
                f.write(content)
            
            image_paths.append(str(img_path))
        
        # Convert images to PDF
        pdf_data = ImageService.images_to_pdf(image_paths)
        
        with open(output_path, "wb") as f:
            f.write(pdf_data)
        
        return {
            "status": "success",
            "file_id": upload_id,
            "image_count": len(files),
            "output_size": output_path.stat().st_size
        }
    except Exception as e:
        cleanup(output_path, *image_paths)
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
    finally:
        cleanup(*image_paths)

@router.post("/pdf-to-text")
async def pdf_to_text(file: UploadFile = File(...)):
    """Extract text from PDF"""
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files allowed")
    
    upload_id = uid()
    input_path = UPLOADS / f"{upload_id}.pdf"
    output_path = OUTPUTS / f"{upload_id}_text.txt"
    
    try:
        # Save uploaded file
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)
        
        # Extract text
        doc = fitz.open(input_path)
        text = ""
        for page_num, page in enumerate(doc):
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page.get_text()
        
        doc.close()
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        
        return {
            "status": "success",
            "file_id": upload_id,
            "filename": file.filename,
            "text_length": len(text)
        }
    except Exception as e:
        cleanup(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {str(e)}")

@router.post("/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    """Convert PDF text into an editable Word document (.docx)."""
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files allowed")

    upload_id = uid()
    input_path = UPLOADS / f"{upload_id}.pdf"
    output_path = OUTPUTS / f"{upload_id}_word.docx"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        pdf = fitz.open(input_path)
        docx = Document()
        docx.add_heading(Path(file.filename).stem or "Converted PDF", level=1)

        for page_index, page in enumerate(pdf, start=1):
            if page_index > 1:
                docx.add_page_break()
            docx.add_paragraph(f"Page {page_index}").style = docx.styles["Heading 2"]
            text = page.get_text("text").strip()
            if text:
                for block in text.split("\n\n"):
                    clean = block.strip()
                    if clean:
                        docx.add_paragraph(clean)
            else:
                docx.add_paragraph("[No selectable text found on this page]")

        page_count = len(pdf)
        pdf.close()
        docx.save(output_path)

        return {
            "status": "success",
            "file_id": upload_id,
            "filename": file.filename,
            "page_count": page_count
        }
    except Exception as e:
        cleanup(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Word conversion failed: {str(e)}")
    finally:
        cleanup(input_path)


@router.get("/download-image/{file_id}")
async def download_images(file_id: str):
    """Download converted images as ZIP"""
    zip_path = OUTPUTS / f"{file_id}_images.zip"
    
    if not zip_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(zip_path, filename=f"images_{file_id}.zip")

@router.get("/download-pdf/{file_id}")
async def download_pdf(file_id: str):
    """Download converted PDF"""
    output_path = OUTPUTS / f"{file_id}_from_images.pdf"
    
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(output_path, filename=f"converted_{file_id}.pdf")

@router.get("/download-text/{file_id}")
async def download_text(file_id: str):
    """Download extracted text"""
    output_path = OUTPUTS / f"{file_id}_text.txt"

    if not output_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(output_path, filename=f"text_{file_id}.txt")

@router.get("/download-word/{file_id}")
async def download_word(file_id: str):
    """Download converted Word document"""
    output_path = OUTPUTS / f"{file_id}_word.docx"

    if not output_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(output_path, filename=f"word_{file_id}.docx")
